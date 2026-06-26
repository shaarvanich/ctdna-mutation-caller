import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in dxas (1/20th of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_runs(p, text):
    """Parses text for inline markdown formatting like **bold** and *italic* and adds to paragraph."""
    # Split text by bold markers (**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold segment
            inner_text = part[2:-2]
            # Check for italic within bold
            subparts = re.split(r'(\*.*?\*)', inner_text)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = p.add_run(subpart[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = p.add_run(subpart)
                    run.bold = True
        else:
            # Check for italic (*)
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = p.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    # Regular text
                    p.add_run(subpart)

def convert_md_to_docx(md_path, docx_path):
    print(f"Reading markdown from: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # ── Page Margins Configuration ──────────────────────────────────────────
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Set Base Font and Styles ──────────────────────────────────────────────
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Dark charcoal text

    # Set Heading 1 font style
    style_h1 = doc.styles['Heading 1']
    h1_font = style_h1.font
    h1_font.name = 'Calibri Light'
    h1_font.size = Pt(20)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)  # Deep Navy

    # Set Heading 2 font style
    style_h2 = doc.styles['Heading 2']
    h2_font = style_h2.font
    h2_font.name = 'Calibri'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)  # Muted Blue

    # Set Heading 3 font style
    style_h3 = doc.styles['Heading 3']
    h3_font = style_h3.font
    h3_font.name = 'Calibri'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0x59, 0x59, 0x59)  # Slate Gray

    i = 0
    in_code_block = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines but add a small space if we want paragraph separation
        if not line:
            # We don't need double paragraphs, docx adds space after automatically
            i += 1
            continue

        # Check for code blocks (e.g. ```mermaid)
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            # Add code line styled as gray block
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Check for Markdown horizontal rule
        if line == '---':
            # Add border or page break or horizontal line
            doc.add_page_break()
            i += 1
            continue

        # ── Heading 1 ──────────────────────────────────────────────────────────
        if line.startswith('# '):
            p = doc.add_paragraph(style='Heading 1')
            add_paragraph_runs(p, line[2:])
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Heading 2 ──────────────────────────────────────────────────────────
        elif line.startswith('## '):
            p = doc.add_paragraph(style='Heading 2')
            add_paragraph_runs(p, line[3:])
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # ── Heading 3 ──────────────────────────────────────────────────────────
        elif line.startswith('### '):
            p = doc.add_paragraph(style='Heading 3')
            add_paragraph_runs(p, line[4:])
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # ── Heading 4 ──────────────────────────────────────────────────────────
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[5:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            i += 1
            continue

        # ── Heading 5 ──────────────────────────────────────────────────────────
        elif line.startswith('##### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[6:])
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
            i += 1
            continue

        # ── Blockquotes / Alerts ──────────────────────────────────────────────
        elif line.startswith('>'):
            # Group contiguous blockquotes
            quote_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                cleaned_quote = lines[i].rstrip()[1:].strip()
                # Skip markdown alert type tags like [!IMPORTANT] in blockquote
                if not (cleaned_quote.startswith('[!') and cleaned_quote.endswith(']')):
                    quote_lines.append(cleaned_quote)
                i += 1
            
            quote_text = " ".join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Format text inside blockquote
            add_paragraph_runs(p, quote_text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # ── Bullet Lists ──────────────────────────────────────────────────────
        elif line.startswith('* ') or line.startswith('- '):
            # Check prefix length
            p = doc.add_paragraph(style='List Bullet')
            add_paragraph_runs(p, line[2:])
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # ── Numbered Lists ────────────────────────────────────────────────────
        elif re.match(r'^\d+\.\s', line):
            prefix_match = re.match(r'^(\d+\.)\s', line)
            prefix = prefix_match.group(0)
            p = doc.add_paragraph(style='List Number')
            add_paragraph_runs(p, line[len(prefix):])
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────────────────
        elif line.startswith('|'):
            # Group contiguous table rows
            table_lines = []
            while i < len(lines) and lines[i].rstrip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            
            # Parse rows
            rows_data = []
            for tl in table_lines:
                # Remove starting and ending pipes
                cells = tl.strip().split('|')
                # Filter out empty first and last split values
                cells = [c.strip() for c in cells if c.strip() != '']
                # Ignore divider rows like | :--- | :---: |
                if cells and all(re.match(r'^:?-+:?$', c) for c in cells):
                    continue
                if cells:
                    rows_data.append(cells)
            
            if not rows_data:
                continue

            num_cols = len(rows_data[0])
            num_rows = len(rows_data)

            # Create Table in docx
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.autofit = True
            
            for r_idx, row_cells in enumerate(rows_data):
                row = table.rows[r_idx]
                for c_idx, cell_value in enumerate(row_cells):
                    if c_idx >= num_cols:
                        break
                    cell = row.cells[c_idx]
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    
                    # Style headers
                    if r_idx == 0:
                        set_cell_background(cell, "1F4E78")  # Dark blue header
                        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                        add_paragraph_runs(p, f"**{cell_value}**")
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text
                            run.font.size = Pt(10)
                    else:
                        # Light shading for zebra stripes
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "F2F2F2")
                        else:
                            set_cell_background(cell, "FFFFFF")
                        set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
                        add_paragraph_runs(p, cell_value)
                        for run in p.runs:
                            run.font.size = Pt(9.5)
            
            doc.add_paragraph()  # Space below table
            continue

        # ── Display Equations ─────────────────────────────────────────────────
        elif line.startswith('$$') and line.endswith('$$'):
            equation = line[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(equation)
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            i += 1
            continue

        # ── Standard Paragraph ────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_paragraph_runs(p, line)
            i += 1

    print(f"Saving docx to: {docx_path}")
    doc.save(docx_path)
    print("Conversion complete!")

if __name__ == "__main__":
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_file = os.path.join(base_dir, "outputs", "reports", "academic_project_report.md")
    docx_file = os.path.join(base_dir, "outputs", "reports", "academic_project_report.docx")
    
    convert_md_to_docx(md_file, docx_file)
